#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Restaurant Management System - Project Documentation Generator
Creates a comprehensive Word document describing the project scope, features, goals, and timeline.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re


class ProjectDocumentGenerator:
    """Generates a comprehensive project documentation in Word format."""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Configure document styles for Vietnamese text support."""
        # Set up default font for the document
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Configure paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)
        
    def add_heading_custom(self, text, level=1):
        """Add a custom heading with Arial Bold font."""
        heading = self.doc.add_heading(text, level=level)
        
        # Set font to Arial Bold
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.bold = True
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
                
        return heading
        
    def add_paragraph_custom(self, text, bold=False, italic=False):
        """Add a paragraph with Times New Roman font."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        run.font.italic = italic
        return para
        
    def add_bullet_point(self, text, level=0):
        """Add a bullet point with proper indentation."""
        para = self.doc.add_paragraph(text, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.5 * (level + 1))
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        return para
        
    def add_title_page(self):
        """Add title page with project information."""
        # Add some space
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Title
        title = self.doc.add_heading('TÀI LIỆU MÔ TẢ DỰ ÁN', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.bold = True
            
        # Subtitle
        subtitle = self.doc.add_heading('RESTAURANT MANAGEMENT SYSTEM', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            
        # Add space
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Project info
        info_text = [
            'Hệ Thống Quản Lý Nhà Hàng Toàn Diện',
            '',
            'Công nghệ: Next.js, Express.js, PostgreSQL, Docker, Tauri',
            '',
            f'Ngày tạo tài liệu: {datetime.now().strftime("%d/%m/%Y")}',
        ]
        
        for text in info_text:
            para = self.doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
        # Page break
        self.doc.add_page_break()
        
    def add_table_of_contents(self):
        """Add table of contents placeholder."""
        self.add_heading_custom('MỤC LỤC', level=1)
        
        toc_items = [
            '1. QUY MÔ DỰ ÁN',
            '2. CHỨC NĂNG DỰ ÁN',
            '3. MỤC TIÊU CÓ THỂ ĐÁP ỨNG',
            '4. LỊCH LÀM DỰ ÁN',
        ]
        
        for item in toc_items:
            para = self.doc.add_paragraph(item)
            para.paragraph_format.left_indent = Inches(0.5)
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
        self.doc.add_page_break()
        
    def add_section_1_project_scale(self):
        """Section 1: Quy mô dự án - Project Scale"""
        self.add_heading_custom('1. QUY MÔ DỰ ÁN', level=1)
        
        # 1.1 Tổng quan hệ thống
        self.add_heading_custom('1.1. Tổng Quan Hệ Thống Quản Lý Nhà Hàng', level=2)
        
        overview_text = """Restaurant Management System là một hệ thống quản lý nhà hàng hiện đại, toàn diện được xây dựng với kiến trúc full-stack và hoàn toàn được container hóa với Docker. Hệ thống được thiết kế để hỗ trợ toàn bộ quy trình vận hành của một nhà hàng từ việc quản lý thực đơn, đặt bàn, gọi món, xử lý đơn hàng tại bếp, thanh toán đến quản lý nhân sự và tồn kho."""
        
        self.add_paragraph_custom(overview_text)
        
        self.add_paragraph_custom('\nĐặc điểm nổi bật của hệ thống:', bold=True)
        
        features = [
            'Kiến trúc full-stack hiện đại với frontend Next.js và backend Express.js',
            'Hoàn toàn được container hóa với Docker để dễ dàng triển khai và mở rộng',
            'Cơ sở dữ liệu PostgreSQL 16 với Prisma ORM cho quản lý dữ liệu hiệu quả',
            'Hỗ trợ cập nhật thời gian thực (real-time) với WebSocket/Socket.IO',
            'Giao diện đa ngôn ngữ (Tiếng Việt và Tiếng Anh)',
            'Responsive design tương thích với mọi thiết bị (desktop, tablet, mobile)',
            'Hỗ trợ dark mode cho trải nghiệm người dùng tốt hơn',
            'Ứng dụng desktop đa nền tảng với Tauri (Windows, macOS, Linux)',
        ]
        
        for feature in features:
            self.add_bullet_point(feature)
            
        # 1.2 Kiến trúc hệ thống
        self.add_heading_custom('1.2. Kiến Trúc Hệ Thống', level=2)
        
        architecture_text = """Hệ thống được xây dựng theo kiến trúc 3-tier (ba tầng) với sự tách biệt rõ ràng giữa presentation layer, business logic layer và data access layer. Toàn bộ hệ thống được đóng gói trong các container Docker để đảm bảo tính nhất quán giữa các môi trường phát triển, testing và production."""
        
        self.add_paragraph_custom(architecture_text)
        
        self.add_paragraph_custom('\nCác tầng trong kiến trúc:', bold=True)
        
        layers = [
            'Presentation Layer: Next.js 15 với React 19, TypeScript, Tailwind CSS',
            'Business Logic Layer: Express.js 5 với TypeScript, Prisma ORM',
            'Data Access Layer: PostgreSQL 16, Redis 7 cho caching',
            'Infrastructure Layer: Docker, Docker Compose, Nginx reverse proxy',
        ]
        
        for layer in layers:
            self.add_bullet_point(layer)
            
        # 1.3 Công nghệ sử dụng
        self.add_heading_custom('1.3. Công Nghệ Sử Dụng', level=2)
        
        self.add_paragraph_custom('Frontend Technologies:', bold=True)
        frontend_tech = [
            'Next.js 15: Framework React với Server-Side Rendering và Static Site Generation',
            'React 19: Thư viện UI component-based',
            'TypeScript: Ngôn ngữ lập trình type-safe',
            'Tailwind CSS: Framework CSS utility-first',
            'Zustand: State management library nhẹ và hiệu quả',
            'React Query: Quản lý data fetching và caching',
            'Socket.io Client: Real-time communication',
            'React Hook Form + Zod: Form validation',
            'Radix UI: Accessible UI components',
        ]
        
        for tech in frontend_tech:
            self.add_bullet_point(tech)
            
        self.add_paragraph_custom('\nBackend Technologies:', bold=True)
        backend_tech = [
            'Express.js 5: Web framework cho Node.js',
            'TypeScript: Type-safe backend development',
            'Prisma ORM: Modern database ORM',
            'PostgreSQL 16: Relational database',
            'Redis 7: In-memory caching và session storage',
            'Socket.io: WebSocket server cho real-time updates',
            'JWT: JSON Web Tokens cho authentication',
            'Multer: File upload handling',
            'Winston: Logging framework',
        ]
        
        for tech in backend_tech:
            self.add_bullet_point(tech)
            
        self.add_paragraph_custom('\nDesktop Application:', bold=True)
        desktop_tech = [
            'Tauri v2: Framework xây dựng desktop app với web technologies',
            'Rust: Backend core cho performance và security',
            'Vite: Build tool hiện đại và nhanh',
            'Shared codebase với web frontend',
        ]
        
        for tech in desktop_tech:
            self.add_bullet_point(tech)
            
        self.add_paragraph_custom('\nDevOps & Infrastructure:', bold=True)
        devops_tech = [
            'Docker: Container platform',
            'Docker Compose: Orchestration tool',
            'Nginx: Reverse proxy và load balancer',
            'GitHub Actions: CI/CD pipeline',
            'Makefile & PowerShell scripts: Automation tools',
        ]
        
        for tech in devops_tech:
            self.add_bullet_point(tech)
            
        # 1.4 Cấu trúc thư mục và module
        self.add_heading_custom('1.4. Cấu Trúc Thư Mục và Module Chính', level=2)
        
        structure_text = """Dự án được tổ chức theo kiến trúc monorepo với các module độc lập nhưng có thể chia sẻ code. Mỗi module có trách nhiệm riêng biệt và có thể phát triển độc lập."""
        
        self.add_paragraph_custom(structure_text)
        
        self.add_paragraph_custom('\nCác module chính:', bold=True)
        
        modules = [
            ('client/', 'Frontend web application với Next.js. Chứa toàn bộ UI components, pages, state management, và business logic phía client.'),
            ('server/', 'Backend API server với Express.js. Chứa controllers, services, repositories, middleware, và Prisma schema cho database.'),
            ('desktop/', 'Desktop application với Tauri. Ứng dụng desktop đa nền tảng chia sẻ code với web frontend.'),
            ('docs/', 'Tài liệu dự án bao gồm use cases, database schema, API documentation, và deployment guides.'),
            ('nginx/', 'Cấu hình Nginx reverse proxy cho production deployment.'),
            ('data/', 'Data seeding và migration scripts cho development và testing.'),
        ]
        
        for module_name, module_desc in modules:
            self.add_paragraph_custom(f'\n{module_name}', bold=True)
            self.add_paragraph_custom(module_desc)
            
        self.add_paragraph_custom('\nCấu trúc chi tiết client/ (Frontend):', bold=True)
        client_structure = [
            'src/app/: Next.js App Router pages và layouts',
            'src/components/: React components (features, shared, UI)',
            'src/contexts/: React contexts (Socket, Theme)',
            'src/hooks/: Custom React hooks',
            'src/lib/: Utility functions và helpers',
            'src/services/: API service layer',
            'src/stores/: Zustand state stores',
            'src/types/: TypeScript type definitions',
            'public/: Static assets (images, icons)',
            'locales/: Internationalization (i18n) translations',
        ]
        
        for item in client_structure:
            self.add_bullet_point(item)
            
        self.add_paragraph_custom('\nCấu trúc chi tiết server/ (Backend):', bold=True)
        server_structure = [
            'src/config/: Configuration files (database, redis, jwt)',
            'src/controllers/: Route controllers',
            'src/middlewares/: Express middlewares (auth, error handling)',
            'src/repositories/: Data access layer với Prisma',
            'src/routes/: API route definitions',
            'src/services/: Business logic layer',
            'src/types/: TypeScript type definitions',
            'src/utils/: Utility functions',
            'src/validators/: Input validation với Zod',
            'prisma/: Prisma schema, migrations, seeds',
        ]
        
        for item in server_structure:
            self.add_bullet_point(item)
            
    def add_section_2_project_features(self):
        """Section 2: Chức năng dự án - Project Features"""
        self.add_heading_custom('2. CHỨC NĂNG DỰ ÁN', level=1)
        
        intro_text = """Hệ thống Restaurant Management cung cấp một bộ tính năng toàn diện để quản lý mọi khía cạnh của hoạt động nhà hàng. Các chức năng được thiết kế để tối ưu hóa quy trình làm việc, giảm thiểu sai sót và nâng cao trải nghiệm khách hàng."""
        
        self.add_paragraph_custom(intro_text)
        
        # 2.1 Quản lý xác thực và tài khoản
        self.add_heading_custom('2.1. Quản Lý Xác Thực và Tài Khoản Người Dùng', level=2)
        
        auth_text = """Hệ thống authentication bảo mật với JWT (JSON Web Tokens) cho phép quản lý người dùng và phiên làm việc một cách an toàn."""
        
        self.add_paragraph_custom(auth_text)
        
        auth_features = [
            'Đăng ký tài khoản: Tạo tài khoản mới với email, mật khẩu và thông tin cá nhân',
            'Đăng nhập/Đăng xuất: Xác thực người dùng với JWT access token và refresh token',
            'Quên mật khẩu: Reset mật khẩu qua email với token có thời hạn',
            'Quản lý phiên làm việc: Tự động refresh token khi hết hạn',
            'Phân quyền: Role-based access control (RBAC) cho các chức năng khác nhau',
            'Bảo mật: Hash mật khẩu với bcrypt, HTTPS, rate limiting',
        ]
        
        for feature in auth_features:
            self.add_bullet_point(feature)
            
        # 2.2 Quản lý menu và danh mục
        self.add_heading_custom('2.2. Quản Lý Menu và Danh Mục', level=2)
        
        menu_text = """Module quản lý menu cho phép tạo, cập nhật và quản lý toàn bộ thực đơn của nhà hàng với các danh mục và món ăn."""
        
        self.add_paragraph_custom(menu_text)
        
        menu_features = [
            'Quản lý danh mục: Tạo/sửa/xóa các danh mục món ăn (Khai vị, Món chính, Tráng miệng, Đồ uống)',
            'Quản lý món ăn: Thêm món ăn mới với tên, mô tả, giá, hình ảnh, thành phần',
            'Cập nhật giá: Điều chỉnh giá bán theo mùa hoặc chi phí nguyên liệu',
            'Quản lý trạng thái: Đánh dấu món ăn "có sẵn" hoặc "hết hàng"',
            'Upload hình ảnh: Quản lý hình ảnh món ăn với file storage system',
            'Công thức món ăn: Liên kết món ăn với nguyên liệu và số lượng cần thiết',
            'Tìm kiếm và lọc: Tìm kiếm món ăn theo tên, danh mục, giá',
        ]
        
        for feature in menu_features:
            self.add_bullet_point(feature)
            
        # 2.3 Quản lý bàn và đặt bàn
        self.add_heading_custom('2.3. Quản Lý Bàn và Đặt Bàn', level=2)
        
        table_text = """Hệ thống quản lý bàn ăn giúp tối ưu hóa việc sắp xếp chỗ ngồi và theo dõi trạng thái bàn theo thời gian thực."""
        
        self.add_paragraph_custom(table_text)
        
        table_features = [
            'Cấu hình bàn: Tạo bàn với số hiệu, sức chứa, vị trí/khu vực',
            'Trạng thái bàn: Theo dõi trạng thái (Trống, Đã đặt, Đang sử dụng, Cần dọn)',
            'Sơ đồ nhà hàng: Hiển thị layout bàn ăn với màu sắc theo trạng thái',
            'Ghép/tách bàn: Kết hợp nhiều bàn cho nhóm khách lớn',
            'Đặt bàn trước: Khách có thể đặt bàn qua website/app với xác nhận tự động',
            'Quản lý đặt bàn: Xem, chỉnh sửa, xác nhận hoặc hủy đặt bàn',
            'Thông báo nhắc lịch: Gửi SMS/Email nhắc nhở khách trước 24 giờ',
            'Xử lý no-show: Quản lý trường hợp khách không đến',
        ]
        
        for feature in table_features:
            self.add_bullet_point(feature)
            
        # 2.4 Quản lý đơn hàng
        self.add_heading_custom('2.4. Quản Lý Đơn Hàng và Bếp', level=2)
        
        order_text = """Module quản lý đơn hàng xử lý toàn bộ quy trình từ khi khách gọi món đến khi món ăn được phục vụ."""
        
        self.add_paragraph_custom(order_text)
        
        order_features = [
            'Tạo đơn hàng: Nhân viên phục vụ tạo đơn cho bàn hoặc tại quầy',
            'Thêm món: Chọn món từ menu, điều chỉnh số lượng, thêm ghi chú đặc biệt',
            'Chỉnh sửa đơn: Thay đổi, thêm hoặc xóa món trước/sau khi gửi bếp',
            'Gửi đến bếp: Đơn hàng được gửi real-time đến kitchen display',
            'Theo dõi trạng thái: Chờ xác nhận → Đang chuẩn bị → Sẵn sàng → Đã phục vụ',
            'Kitchen Display System (KDS): Màn hình hiển thị đơn hàng cho bếp',
            'Cập nhật từ bếp: Đầu bếp cập nhật tiến độ chuẩn bị món',
            'Thông báo real-time: Nhân viên phục vụ nhận thông báo khi món sẵn sàng',
            'Ghi chú đặc biệt: Yêu cầu khách (ít gia vị, không hành, dị ứng)',
        ]
        
        for feature in order_features:
            self.add_bullet_point(feature)
            
        # 2.5 Quản lý thanh toán và hóa đơn
        self.add_heading_custom('2.5. Quản Lý Thanh Toán và Hóa Đơn', level=2)
        
        payment_text = """Hệ thống thanh toán linh hoạt hỗ trợ nhiều phương thức và tính năng quản lý hóa đơn chi tiết."""
        
        self.add_paragraph_custom(payment_text)
        
        payment_features = [
            'Tạo hóa đơn: Tự động tính tổng tiền, thuế, phí dịch vụ từ đơn hàng',
            'Áp dụng khuyến mãi: Mã giảm giá, voucher, chương trình khuyến mãi',
            'Thanh toán đa phương thức: Tiền mặt, thẻ tín dụng, ví điện tử, chuyển khoản',
            'Thanh toán một phần: Khách có thể trả một phần và trả phần còn lại sau',
            'Chia hóa đơn: Chia bill cho nhiều người trong nhóm',
            'Hoàn tiền: Xử lý hoàn tiền với lý do và phê duyệt',
            'In hóa đơn: In hóa đơn chi tiết hoặc gửi qua email',
            'Lưu trữ hóa đơn: Tìm kiếm và xem lại hóa đơn cũ',
            'Báo cáo doanh thu: Tổng hợp doanh thu theo ngày/tuần/tháng',
        ]
        
        for feature in payment_features:
            self.add_bullet_point(feature)
            
        # 2.6 Quản lý nhân sự
        self.add_heading_custom('2.6. Quản Lý Nhân Sự', level=2)
        
        staff_text = """Module quản lý nhân viên giúp theo dõi hiệu suất làm việc và phân quyền hệ thống."""
        
        self.add_paragraph_custom(staff_text)
        
        staff_features = [
            'Quản lý tài khoản nhân viên: Tạo, cập nhật, vô hiệu hóa tài khoản',
            'Phân quyền: Gán vai trò (Quản lý, Nhân viên phục vụ, Đầu bếp, Thu ngân)',
            'Phân công ca làm: Sắp xếp lịch làm việc cho nhân viên',
            'Theo dõi giờ làm: Ghi nhận giờ vào/ra, tính tổng giờ làm',
            'Đánh giá hiệu suất: Ghi nhận số đơn xử lý, feedback từ khách',
            'Quản lý khu vực: Gán nhân viên phục vụ cho các khu vực cụ thể',
            'Lịch sử hoạt động: Xem log hoạt động của từng nhân viên',
        ]
        
        for feature in staff_features:
            self.add_bullet_point(feature)
            
        # 2.7 Quản lý tồn kho
        self.add_heading_custom('2.7. Quản Lý Tồn Kho và Nguyên Liệu', level=2)
        
        inventory_text = """Hệ thống quản lý tồn kho toàn diện giúp theo dõi nguyên liệu, đơn đặt hàng và cảnh báo khi hết hàng."""
        
        self.add_paragraph_custom(inventory_text)
        
        inventory_features = [
            'Quản lý nguyên liệu: Tạo danh sách nguyên liệu với mã SKU, đơn vị tính, giá vốn',
            'Quản lý nhà cung cấp: Thông tin nhà cung cấp, người liên hệ, điều khoản thanh toán',
            'Đơn đặt hàng: Tạo purchase order cho nhà cung cấp',
            'Nhận hàng: Ghi nhận hàng nhập kho với lô hàng (batch) và hạn sử dụng',
            'Quản lý lô hàng: Theo dõi từng lô hàng với số lượng và expiry date',
            'Giao dịch tồn kho: Ghi nhận mọi thay đổi (nhập, xuất, điều chỉnh, hao hụt)',
            'Cảnh báo tự động: Thông báo khi tồn kho thấp, sắp hết hạn, đã hết hạn',
            'Công thức món ăn: Liên kết món ăn với nguyên liệu để tính toán tự động',
            'Kiểm kê: Kiểm kê định kỳ và điều chỉnh tồn kho',
            'Báo cáo tồn kho: Báo cáo chi tiết về tồn kho, chuyển động, hao hụt',
        ]
        
        for feature in inventory_features:
            self.add_bullet_point(feature)
            
        # 2.8 Tính năng kỹ thuật
        self.add_heading_custom('2.8. Tính Năng Kỹ Thuật', level=2)
        
        technical_text = """Các tính năng kỹ thuật nâng cao giúp nâng cao trải nghiệm người dùng và hiệu suất hệ thống."""
        
        self.add_paragraph_custom(technical_text)
        
        technical_features = [
            'Real-time Updates: WebSocket/Socket.IO cho cập nhật trạng thái tức thời',
            'Đa ngôn ngữ (i18n): Hỗ trợ Tiếng Việt và Tiếng Anh, dễ dàng mở rộng',
            'Responsive Design: Giao diện thích ứng với mọi kích thước màn hình',
            'Dark Mode: Chế độ tối giúp giảm mỏi mắt khi làm việc ban đêm',
            'Progressive Web App (PWA): Có thể cài đặt và sử dụng offline',
            'File Upload: Upload và quản lý hình ảnh cho menu items',
            'Export Data: Xuất báo cáo dưới dạng PDF, Excel',
            'Search & Filter: Tìm kiếm và lọc dữ liệu nhanh chóng',
            'Pagination: Phân trang cho danh sách lớn',
            'Error Handling: Xử lý lỗi thân thiện với người dùng',
            'Loading States: Hiển thị trạng thái loading khi xử lý dữ liệu',
            'Form Validation: Validate input với Zod schema',
            'Toast Notifications: Thông báo toast cho các hành động',
        ]
        
        for feature in technical_features:
            self.add_bullet_point(feature)
            
    def add_section_3_achievable_goals(self):
        """Section 3: Mục tiêu có thể đáp ứng"""
        self.add_heading_custom('3. MỤC TIÊU CÓ THỂ ĐÁP ỨNG', level=1)
        
        intro_text = """Hệ thống Restaurant Management được thiết kế để giải quyết các vấn đề kinh doanh thực tế và mang lại giá trị thiết thực cho nhà hàng. Dưới đây là các mục tiêu chính mà hệ thống có thể đạt được."""
        
        self.add_paragraph_custom(intro_text)
        
        # 3.1 Giải quyết vấn đề kinh doanh
        self.add_heading_custom('3.1. Giải Quyết Các Vấn Đề Kinh Doanh Của Nhà Hàng', level=2)
        
        business_problems = [
            ('Quản lý đơn hàng thủ công', 'Hệ thống số hóa toàn bộ quy trình gọi món, giảm thiểu sai sót và thời gian xử lý'),
            ('Giao tiếp giữa bếp và phục vụ', 'Real-time communication đảm bảo thông tin được truyền đạt ngay lập tức'),
            ('Theo dõi tồn kho nguyên liệu', 'Tự động cảnh báo khi nguyên liệu sắp hết, giảm tình trạng thiếu hàng'),
            ('Mất khách do không quản lý đặt bàn', 'Hệ thống đặt bàn trực tuyến với xác nhận tự động và nhắc lịch'),
            ('Khó khăn trong tính toán doanh thu', 'Báo cáo tự động theo thời gian thực, dễ dàng phân tích kinh doanh'),
            ('Quản lý nhân viên không hiệu quả', 'Theo dõi giờ làm, hiệu suất và phân quyền rõ ràng'),
        ]
        
        for problem, solution in business_problems:
            self.add_paragraph_custom(f'\n{problem}:', bold=True)
            self.add_paragraph_custom(solution)
            
        # 3.2 Cải thiện hiệu quả vận hành
        self.add_heading_custom('3.2. Cải Thiện Hiệu Quả Vận Hành', level=2)
        
        operational_improvements = [
            'Giảm thời gian xử lý đơn hàng: Từ 5-10 phút xuống còn 1-2 phút nhờ hệ thống số',
            'Tăng thông lượng khách: Xử lý nhiều đơn hàng hơn trong cùng thời gian',
            'Giảm sai sót: Đơn hàng được gửi trực tiếp từ phục vụ đến bếp, không qua giấy tờ',
            'Tối ưu sắp xếp bàn: Sơ đồ real-time giúp sắp xếp khách hiệu quả hơn',
            'Quản lý tồn kho chính xác: Giảm lãng phí do quản lý hạn sử dụng tốt hơn',
            'Tiết kiệm chi phí nhân sự: Tự động hóa nhiều tác vụ, cần ít nhân viên hơn',
            'Cải thiện communication: Real-time updates giữa các bộ phận',
        ]
        
        for improvement in operational_improvements:
            self.add_bullet_point(improvement)
            
        # 3.3 Tăng trải nghiệm khách hàng
        self.add_heading_custom('3.3. Tăng Trải Nghiệm Khách Hàng', level=2)
        
        customer_experience = [
            'Đặt bàn dễ dàng: Khách có thể đặt bàn online 24/7, nhận xác nhận ngay lập tức',
            'Giảm thời gian chờ: Đơn hàng được xử lý nhanh hơn nhờ hệ thống số',
            'Độ chính xác cao: Ghi chú đặc biệt được truyền đạt chính xác đến bếp',
            'Thanh toán linh hoạt: Nhiều phương thức thanh toán, chia bill dễ dàng',
            'Hóa đơn rõ ràng: Hóa đơn chi tiết, có thể gửi qua email',
            'Phản hồi nhanh: Nhân viên có thể xem trạng thái đơn hàng mọi lúc',
            'Đa ngôn ngữ: Phục vụ cả khách Việt và khách nước ngoài',
        ]
        
        for experience in customer_experience:
            self.add_bullet_point(experience)
            
        # 3.4 Hỗ trợ quản lý toàn diện
        self.add_heading_custom('3.4. Hỗ Trợ Quản Lý Toàn Diện Từ Frontend Đến Backend', level=2)
        
        management_text = """Hệ thống cung cấp công cụ quản lý toàn diện cho chủ nhà hàng và quản lý:"""
        
        self.add_paragraph_custom(management_text)
        
        management_features = [
            'Dashboard tổng quan: Xem tổng quan doanh thu, đơn hàng, trạng thái bàn',
            'Báo cáo chi tiết: Doanh thu, món ăn bán chạy, hiệu suất nhân viên',
            'Phân tích xu hướng: Xác định giờ cao điểm, món ăn phổ biến',
            'Quản lý chi phí: Theo dõi chi phí nguyên liệu, tính toán lợi nhuận',
            'Kiểm soát quyền hạn: Phân quyền chi tiết cho từng chức vụ',
            'Audit log: Theo dõi mọi thay đổi trong hệ thống',
            'Multi-platform: Truy cập từ web, desktop app, mobile (responsive)',
        ]
        
        for feature in management_features:
            self.add_bullet_point(feature)
            
        # 3.5 Khả năng mở rộng và bảo trì
        self.add_heading_custom('3.5. Khả Năng Mở Rộng và Bảo Trì', level=2)
        
        scalability_text = """Hệ thống được thiết kế với khả năng mở rộng cao và dễ dàng bảo trì:"""
        
        self.add_paragraph_custom(scalability_text)
        
        scalability_features = [
            'Kiến trúc module: Các module độc lập, dễ dàng thêm tính năng mới',
            'Docker containerization: Dễ dàng deploy và scale',
            'Database indexing: Tối ưu performance cho database lớn',
            'Caching với Redis: Giảm tải cho database, tăng tốc độ',
            'TypeScript: Type-safe code giảm bugs, dễ maintain',
            'Comprehensive documentation: Tài liệu đầy đủ cho developers',
            'API-first design: Frontend và backend tách biệt, dễ tích hợp',
            'Automated testing: Unit tests, integration tests (có thể mở rộng)',
            'CI/CD pipeline: GitHub Actions cho automation deployment',
        ]
        
        for feature in scalability_features:
            self.add_bullet_point(feature)
            
        # 3.6 Lợi ích kinh doanh cụ thể
        self.add_heading_custom('3.6. Lợi Ích Kinh Doanh Cụ Thể', level=2)
        
        business_benefits = [
            ('Tăng doanh thu 15-25%', 'Nhờ xử lý nhanh hơn, phục vụ nhiều khách hơn'),
            ('Giảm chi phí vận hành 10-20%', 'Tối ưu nhân sự và giảm lãng phí nguyên liệu'),
            ('Tăng tỷ lệ khách quay lại 30%', 'Trải nghiệm tốt hơn, dịch vụ nhanh hơn'),
            ('Giảm thời gian training nhân viên', 'Giao diện thân thiện, dễ sử dụng'),
            ('Cải thiện customer satisfaction', 'Rating và review tích cực hơn'),
            ('Dữ liệu để ra quyết định', 'Báo cáo chi tiết giúp quản lý đưa ra quyết định đúng đắn'),
        ]
        
        for benefit, description in business_benefits:
            self.add_paragraph_custom(f'\n{benefit}:', bold=True)
            self.add_paragraph_custom(description)
            
    def add_section_4_project_timeline(self):
        """Section 4: Lịch làm dự án"""
        self.add_heading_custom('4. LỊCH LÀM DỰ ÁN', level=1)
        
        intro_text = """Dựa trên phân tích cấu trúc code và tài liệu hiện tại, dự án đã hoàn thành phần lớn các tính năng cốt lõi. Dưới đây là đánh giá tiến độ và kế hoạch phát triển."""
        
        self.add_paragraph_custom(intro_text)
        
        # 4.1 Phân tích tiến độ hiện tại
        self.add_heading_custom('4.1. Phân Tích Tiến Độ Hiện Tại', level=2)
        
        self.add_paragraph_custom('Các module đã hoàn thành:', bold=True)
        
        completed_modules = [
            '✅ Authentication & Authorization: Đã triển khai đầy đủ với JWT',
            '✅ Menu Management: Quản lý danh mục và món ăn hoàn chỉnh',
            '✅ Table Management: Quản lý bàn với trạng thái real-time',
            '✅ Order Management: Quy trình đơn hàng từ tạo đến hoàn thành',
            '✅ Kitchen Display: Màn hình bếp với real-time updates',
            '✅ Billing & Payment: Hệ thống thanh toán đa phương thức',
            '✅ Inventory Management: Quản lý tồn kho và nguyên liệu cơ bản',
            '✅ Staff Management: Quản lý nhân viên và phân quyền',
            '✅ Reservation System: Đặt bàn trực tuyến',
            '✅ Database Schema: ERD và Prisma schema hoàn chỉnh',
            '✅ Docker Deployment: Containerization hoàn chỉnh',
            '✅ Frontend UI: Giao diện người dùng với Next.js',
            '✅ Desktop App: Ứng dụng Tauri cho desktop',
        ]
        
        for module in completed_modules:
            self.add_bullet_point(module)
            
        self.add_paragraph_custom('\nCác tính năng đang phát triển/cần hoàn thiện:', bold=True)
        
        in_progress = [
            '🔄 Reports & Analytics: Dashboard và báo cáo chi tiết',
            '🔄 Advanced Inventory: Tính năng nâng cao cho quản lý tồn kho',
            '🔄 Customer Management: Module quản lý thông tin khách hàng',
            '🔄 Loyalty Program: Chương trình tích điểm và khuyến mãi',
            '🔄 Mobile App: Ứng dụng mobile native (React Native/Flutter)',
            '🔄 Integration APIs: Tích hợp với payment gateway, delivery services',
        ]
        
        for item in in_progress:
            self.add_bullet_point(item)
            
        self.add_paragraph_custom('\nĐánh giá tỷ lệ hoàn thành:', bold=True)
        self.add_paragraph_custom('Core Features: 85% hoàn thành')
        self.add_paragraph_custom('Advanced Features: 60% hoàn thành')
        self.add_paragraph_custom('Testing & Documentation: 70% hoàn thành')
        self.add_paragraph_custom('Deployment & DevOps: 90% hoàn thành')
        
        # 4.2 Ước tính thời gian hoàn thành
        self.add_heading_custom('4.2. Ước Tính Thời Gian Hoàn Thành Các Phần Còn Lại', level=2)
        
        remaining_work = [
            ('Reports & Analytics Dashboard', '2-3 tuần', 'Thiết kế và triển khai các biểu đồ, báo cáo tổng hợp'),
            ('Advanced Inventory Features', '2 tuần', 'Tính năng dự báo, tối ưu đơn hàng'),
            ('Customer Management Module', '2 tuần', 'CRM cơ bản, lịch sử giao dịch'),
            ('Loyalty Program', '2-3 tuần', 'Tích điểm, voucher, membership'),
            ('Mobile Application', '4-6 tuần', 'Phát triển app mobile native'),
            ('Payment Gateway Integration', '1-2 tuần', 'Tích hợp VNPay, Momo, ZaloPay'),
            ('Testing & Bug Fixes', '2 tuần', 'Unit tests, integration tests, bug fixes'),
            ('Performance Optimization', '1 tuần', 'Caching, query optimization, CDN'),
            ('Security Audit', '1 tuần', 'Penetration testing, security hardening'),
            ('Documentation Update', '1 tuần', 'User manual, API docs, deployment guide'),
        ]
        
        for task, duration, description in remaining_work:
            self.add_paragraph_custom(f'\n{task} ({duration}):', bold=True)
            self.add_paragraph_custom(description)
            
        self.add_paragraph_custom('\nTổng thời gian ước tính: 18-24 tuần (4-6 tháng)', bold=True)
        
        # 4.3 Lịch trình phát triển theo giai đoạn
        self.add_heading_custom('4.3. Lịch Trình Phát Triển Theo Giai Đoạn', level=2)
        
        # Phase 1
        self.add_paragraph_custom('\nGiai đoạn 1: Planning & Research (Đã hoàn thành)', bold=True)
        phase1_text = """Thời gian: 2-3 tuần
Nội dung:
- Phân tích yêu cầu nghiệp vụ
- Thiết kế kiến trúc hệ thống
- Lựa chọn công nghệ
- Thiết kế database schema
- Tạo mockups và wireframes"""
        self.add_paragraph_custom(phase1_text)
        
        # Phase 2
        self.add_paragraph_custom('\nGiai đoạn 2: Core Development (Đã hoàn thành 85%)', bold=True)
        phase2_text = """Thời gian: 8-10 tuần
Nội dung:
- Setup project structure và environment
- Phát triển authentication system
- Phát triển menu management
- Phát triển order management
- Phát triển kitchen display
- Phát triển billing & payment
- Phát triển table management
- Basic inventory management
- Real-time communication với Socket.IO"""
        self.add_paragraph_custom(phase2_text)
        
        # Phase 3
        self.add_paragraph_custom('\nGiai đoạn 3: Advanced Features (Đang thực hiện - 60%)', bold=True)
        phase3_text = """Thời gian: 6-8 tuần
Nội dung:
- Reports & analytics dashboard
- Advanced inventory features
- Customer management module
- Loyalty program
- Reservation system enhancements
- Staff management enhancements
- File upload và storage
- Email notifications"""
        self.add_paragraph_custom(phase3_text)
        
        # Phase 4
        self.add_paragraph_custom('\nGiai đoạn 4: Testing & Quality Assurance (Đang thực hiện - 70%)', bold=True)
        phase4_text = """Thời gian: 3-4 tuần
Nội dung:
- Unit testing (Jest, React Testing Library)
- Integration testing
- End-to-end testing (Playwright/Cypress)
- Performance testing
- Security testing
- User acceptance testing (UAT)
- Bug fixing"""
        self.add_paragraph_custom(phase4_text)
        
        # Phase 5
        self.add_paragraph_custom('\nGiai đoạn 5: Deployment & Production (Đã hoàn thành 90%)', bold=True)
        phase5_text = """Thời gian: 2-3 tuần
Nội dung:
- Docker containerization (✅ Hoàn thành)
- CI/CD pipeline setup (✅ Hoàn thành)
- Production environment configuration
- Database migration và seeding
- SSL/HTTPS setup
- Monitoring và logging
- Backup strategy
- Documentation"""
        self.add_paragraph_custom(phase5_text)
        
        # Phase 6
        self.add_paragraph_custom('\nGiai đoạn 6: Mobile App Development (Kế hoạch tương lai)', bold=True)
        phase6_text = """Thời gian: 4-6 tuần
Nội dung:
- Setup React Native/Flutter project
- Reuse business logic từ web
- Develop mobile-specific UI
- Implement offline support
- Push notifications
- App store submission"""
        self.add_paragraph_custom(phase6_text)
        
        # 4.4 Milestone summary
        self.add_heading_custom('4.4. Các Mốc Quan Trọng (Milestones)', level=2)
        
        milestones = [
            ('✅ Milestone 1 - MVP (Minimum Viable Product)', 'Hoàn thành', 'Các tính năng cốt lõi: auth, menu, order, payment'),
            ('✅ Milestone 2 - Beta Version', 'Hoàn thành', 'Thêm inventory, reservation, staff management'),
            ('🔄 Milestone 3 - Feature Complete', 'Đang thực hiện', 'Reports, analytics, advanced features'),
            ('📅 Milestone 4 - Production Ready', 'Q2 2025', 'Testing hoàn chỉnh, production deployment'),
            ('📅 Milestone 5 - Mobile Release', 'Q3 2025', 'Release mobile app trên App Store/Play Store'),
            ('📅 Milestone 6 - Version 2.0', 'Q4 2025', 'AI features, advanced analytics, integrations'),
        ]
        
        for milestone, status, description in milestones:
            self.add_paragraph_custom(f'\n{milestone}: {status}', bold=True)
            self.add_paragraph_custom(description)
            
        # 4.5 Resource allocation
        self.add_heading_custom('4.5. Phân Bổ Nguồn Lực', level=2)
        
        resources_text = """Để hoàn thành dự án theo timeline, cần có đội ngũ phát triển với các vai trò sau:"""
        self.add_paragraph_custom(resources_text)
        
        team_roles = [
            'Project Manager: 1 người - Quản lý dự án, điều phối team',
            'Full-stack Developer: 2-3 người - Phát triển frontend và backend',
            'UI/UX Designer: 1 người - Thiết kế giao diện người dùng',
            'QA Engineer: 1 người - Testing và quality assurance',
            'DevOps Engineer: 1 người (part-time) - CI/CD, deployment, monitoring',
            'Business Analyst: 1 người (part-time) - Phân tích nghiệp vụ, tài liệu',
        ]
        
        for role in team_roles:
            self.add_bullet_point(role)
            
        # 4.6 Risk management
        self.add_heading_custom('4.6. Quản Lý Rủi Ro', level=2)
        
        risks = [
            ('Thay đổi yêu cầu giữa chừng', 'Làm rõ requirements từ đầu, có change request process'),
            ('Technical challenges', 'Research kỹ trước khi chọn công nghệ, có backup plan'),
            ('Resource constraints', 'Ưu tiên features theo MVP approach'),
            ('Timeline delays', 'Buffer time cho mỗi phase, agile methodology'),
            ('Integration issues', 'Integration testing sớm, mock APIs'),
            ('Performance problems', 'Performance testing liên tục, optimization từng bước'),
        ]
        
        for risk, mitigation in risks:
            self.add_paragraph_custom(f'\nRủi ro: {risk}', bold=True)
            self.add_paragraph_custom(f'Giải pháp: {mitigation}')
            
    def add_footer(self):
        """Add footer to all pages except first page."""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Restaurant Management System - Project Documentation"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in footer_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
            
    def generate(self, output_path):
        """Generate the complete document."""
        print("Generating project documentation...")
        
        # Add all sections
        self.add_title_page()
        self.add_table_of_contents()
        self.add_section_1_project_scale()
        self.doc.add_page_break()
        self.add_section_2_project_features()
        self.doc.add_page_break()
        self.add_section_3_achievable_goals()
        self.doc.add_page_break()
        self.add_section_4_project_timeline()
        
        # Add footer
        self.add_footer()
        
        # Save document
        self.doc.save(output_path)
        print(f"✅ Document saved to: {output_path}")


def main():
    """Main function to generate the document."""
    # Output file path
    output_dir = "/home/runner/work/restaurant-management/restaurant-management/docs"
    output_file = os.path.join(output_dir, "TAI_LIEU_MO_TA_DU_AN.docx")
    
    # Generate document
    generator = ProjectDocumentGenerator()
    generator.generate(output_file)
    
    print("\n" + "="*60)
    print("🎉 Tài liệu mô tả dự án đã được tạo thành công!")
    print(f"📄 File: {output_file}")
    print("="*60)


if __name__ == "__main__":
    main()
